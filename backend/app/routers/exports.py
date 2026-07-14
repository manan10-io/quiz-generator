"""
exports.py — FastAPI router for /export/* endpoints.
Generates files in PDF, DOCX, Excel, CSV, JSON, TXT, Moodle XML,
Quizizz CSV, and Kahoot CSV formats from a project's questions.
"""
from __future__ import annotations

import csv
import io
import json
import logging
import os
import tempfile
import uuid
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Literal

from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import Response
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.auth.dependencies import get_current_user
from app.models.models import ExportRecord, Project, Question, User
from app.schemas.schemas import ExportRecordOut, ExportHistoryList, DownloadUrlResponse
from app.utils.signed_url import generate_download_token, verify_download_token

logger = logging.getLogger(__name__)
router = APIRouter()

ExportFormat = Literal["pdf", "docx", "excel", "csv", "json", "txt", "moodle", "quizizz", "kahoot"]

EXPORT_DIR = Path(tempfile.gettempdir()) / "quizgen_exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

EXPORT_EXPIRY_HOURS = 24

FORMAT_MIME_TYPES: dict[str, str] = {
    "csv": "text/csv",
    "json": "application/json",
    "txt": "text/plain",
    "moodle": "application/xml",
    "quizizz": "text/csv",
    "kahoot": "text/csv",
    "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _as_aware_utc(dt: datetime) -> datetime:
    """
    SQLite doesn't preserve tzinfo on DateTime(timezone=True) columns, so
    values read back from the DB come back naive even though they were
    stored as UTC — comparing them against datetime.now(timezone.utc)
    directly raises TypeError. Values were always written as UTC, so a
    naive value is assumed to already be UTC.
    """
    return dt if dt.tzinfo is not None else dt.replace(tzinfo=timezone.utc)


def _storage_path(record_id: str, file_name: str) -> Path:
    """
    On-disk path for a given export record.

    Prefixing with the record's UUID prevents filename collisions between
    different users' (or different projects') exports that happen to share
    a display name — e.g. two teachers both naming a project "Quiz 1" and
    exporting to CSV would otherwise silently overwrite each other's files
    on disk since the old code used the bare file_name as the storage key.
    """
    return EXPORT_DIR / f"{record_id}__{file_name}"


# ─── GET /export/history ──────────────────────────────────────────────────────

@router.get("/history", response_model=ExportHistoryList)
async def get_export_history(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    base_q = select(ExportRecord).where(ExportRecord.user_id == current_user.id)

    q = (
        base_q.order_by(ExportRecord.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )
    result = await db.execute(q)
    records = result.scalars().all()
    return ExportHistoryList(items=[ExportRecordOut.model_validate(r) for r in records])


# ─── GET /export/download/{record_id} ────────────────────────────────────────

@router.get("/download/{record_id}", response_model=DownloadUrlResponse)
async def get_download_url(
    record_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Mints a fresh signed download URL for an existing export record.
    Re-minting on each call (rather than reusing a stored URL) means the
    link's lifetime always reflects the record's actual expires_at, even
    if the original generation-time link was already shared or expired.
    """
    result = await db.execute(
        select(ExportRecord).where(
            ExportRecord.id == record_id,
            ExportRecord.user_id == current_user.id,
        )
    )
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Export record not found")

    if record.expires_at and _as_aware_utc(record.expires_at) < datetime.now(timezone.utc):
        raise HTTPException(
            status.HTTP_410_GONE,
            "This export has expired. Please generate it again.",
        )

    token = generate_download_token(record.id, record.expires_at)
    signed_url = f"/api/v1/export/file/{token}"

    return DownloadUrlResponse(
        url=signed_url,
        file_name=record.file_name,
        expires_at=record.expires_at,
    )


# ─── GET /export/file/{token} ─────────────────────────────────────────────────

@router.get("/file/{token}")
async def download_file(
    token: str,
    db: AsyncSession = Depends(get_db),
):
    """
    Serves the actual export file bytes for a valid, unexpired signed token.

    Deliberately unauthenticated (no get_current_user dependency) since the
    token itself is the credential — this lets the frontend hand the URL
    directly to a browser `<a download>` or window.open() without needing
    to attach an Authorization header, which plain links can't do anyway.
    """
    record_id = verify_download_token(token)
    if not record_id:
        raise HTTPException(
            status.HTTP_401_UNAUTHORIZED,
            "This download link is invalid or has expired.",
        )

    result = await db.execute(select(ExportRecord).where(ExportRecord.id == record_id))
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Export record not found")

    # Defense in depth: even if the token's own embedded expiry somehow
    # passed verification, also honour the DB record's expiry directly.
    if record.expires_at and _as_aware_utc(record.expires_at) < datetime.now(timezone.utc):
        raise HTTPException(status.HTTP_410_GONE, "This export has expired.")

    file_path = _storage_path(record.id, record.file_name)
    if not file_path.exists():
        raise HTTPException(
            status.HTTP_404_NOT_FOUND,
            "Export file is no longer available on the server.",
        )

    mime_type = FORMAT_MIME_TYPES.get(record.format, "application/octet-stream")
    file_bytes = file_path.read_bytes()

    return Response(
        content=file_bytes,
        media_type=mime_type,
        headers={
            "Content-Disposition": f'attachment; filename="{record.file_name}"'
        },
    )


# ─── POST /export/{project_id}/{format} ──────────────────────────────────────

@router.post("/{project_id}/{format}", response_model=ExportRecordOut)
async def export_project(
    project_id: str,
    format: ExportFormat,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Verify ownership
    proj_r = await db.execute(
        select(Project).where(
            Project.id == project_id,
            Project.user_id == current_user.id,
        )
    )
    project = proj_r.scalar_one_or_none()
    if not project:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    # Fetch questions
    q_r = await db.execute(
        select(Question)
        .where(Question.project_id == project_id)
        .order_by(Question.position)
    )
    questions = q_r.scalars().all()

    if not questions:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "Project has no questions to export",
        )

    try:
        file_bytes, file_name, mime_type = _generate_export(
            questions, project.name, format
        )
    except Exception as e:
        logger.error("Export failed for project %s format %s: %s", project_id, format, e)
        raise HTTPException(
            status.HTTP_500_INTERNAL_SERVER_ERROR,
            f"Export generation failed: {e}",
        )

    # Generate the record ID up front so we can use it as part of the
    # on-disk storage path — this is what makes the path collision-safe
    # across different users/projects that happen to share a display name.
    record_id = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(hours=EXPORT_EXPIRY_HOURS)

    out_path = _storage_path(record_id, file_name)
    out_path.write_bytes(file_bytes)

    download_token = generate_download_token(record_id, expires_at)
    signed_url = f"/api/v1/export/file/{download_token}"

    record = ExportRecord(
        id=record_id,
        project_id=project_id,
        user_id=current_user.id,
        format=format,
        file_name=file_name,
        file_size=len(file_bytes),
        download_url=signed_url,
        expires_at=expires_at,
    )
    db.add(record)
    await db.flush()
    return ExportRecordOut.model_validate(record)


# ─── Export generators ────────────────────────────────────────────────────────

def _generate_export(
    questions: list[Question],
    project_name: str,
    format: str,
) -> tuple[bytes, str, str]:
    """Dispatch to the appropriate generator and return (bytes, filename, mime)."""

    safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in project_name)[:50]

    dispatchers = {
        "csv": _export_csv,
        "json": _export_json,
        "txt": _export_txt,
        "moodle": _export_moodle,
        "quizizz": _export_quizizz,
        "kahoot": _export_kahoot,
        "excel": _export_excel,
        "pdf": _export_pdf,
        "docx": _export_docx,
    }

    fn = dispatchers.get(format)
    if not fn:
        raise ValueError(f"Unknown export format: {format}")

    return fn(questions, safe_name)


def _export_csv(questions, name) -> tuple[bytes, str, str]:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["#", "Question", "A", "B", "C", "D", "Answer", "Marks", "Topic", "Difficulty", "Explanation"])
    for i, q in enumerate(questions, 1):
        writer.writerow([
            i, q.question_text,
            q.option_a, q.option_b, q.option_c, q.option_d,
            q.correct_answer or "",
            float(q.marks),
            q.topic or "",
            q.difficulty or "",
            q.explanation or "",
        ])
    return buf.getvalue().encode("utf-8"), f"{name}.csv", "text/csv"


def _export_json(questions, name) -> tuple[bytes, str, str]:
    data = [
        {
            "number": i,
            "question": q.question_text,
            "options": {
                "A": q.option_a, "B": q.option_b,
                "C": q.option_c, "D": q.option_d,
                **({"E": q.option_e} if q.option_e else {}),
            },
            "answer": q.correct_answer,
            "marks": float(q.marks),
            "negative_marks": float(q.negative_marks),
            "topic": q.topic,
            "difficulty": q.difficulty,
            "explanation": q.explanation,
        }
        for i, q in enumerate(questions, 1)
    ]
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8"), f"{name}.json", "application/json"


def _export_txt(questions, name) -> tuple[bytes, str, str]:
    lines: list[str] = []
    for i, q in enumerate(questions, 1):
        lines.append(f"{i}. {q.question_text}")
        for letter, text in [("A", q.option_a), ("B", q.option_b), ("C", q.option_c), ("D", q.option_d)]:
            if text:
                lines.append(f"   {letter}. {text}")
        if q.correct_answer:
            lines.append(f"   Answer: {q.correct_answer}")
        if q.explanation:
            lines.append(f"   Explanation: {q.explanation}")
        lines.append("")
    return "\n".join(lines).encode("utf-8"), f"{name}.txt", "text/plain"


def _export_moodle(questions, name) -> tuple[bytes, str, str]:
    """Moodle XML GIFT-compatible format."""
    root = ET.Element("quiz")
    for i, q in enumerate(questions, 1):
        question_el = ET.SubElement(root, "question", type="multichoice")

        name_el = ET.SubElement(question_el, "name")
        ET.SubElement(name_el, "text").text = f"Question {i}"

        qtext_el = ET.SubElement(question_el, "questiontext", format="html")
        ET.SubElement(qtext_el, "text").text = q.question_text

        ET.SubElement(question_el, "defaultgrade").text = str(float(q.marks))
        ET.SubElement(question_el, "single").text = "true"

        for letter, text in [("A", q.option_a), ("B", q.option_b), ("C", q.option_c), ("D", q.option_d)]:
            if not text:
                continue
            is_correct = q.correct_answer == letter
            answer_el = ET.SubElement(
                question_el, "answer",
                fraction="100" if is_correct else "0",
                format="html",
            )
            ET.SubElement(answer_el, "text").text = text
            if is_correct and q.explanation:
                feedback = ET.SubElement(answer_el, "feedback", format="html")
                ET.SubElement(feedback, "text").text = q.explanation

    xml_bytes = ET.tostring(root, encoding="unicode", xml_declaration=False)
    full_xml = f'<?xml version="1.0" encoding="UTF-8"?>\n{xml_bytes}'
    return full_xml.encode("utf-8"), f"{name}_moodle.xml", "application/xml"


def _export_quizizz(questions, name) -> tuple[bytes, str, str]:
    """Quizizz CSV import format."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Question Text", "Option 1", "Option 2", "Option 3", "Option 4",
                     "Correct Answer", "Time Limit (sec)", "Image Link"])
    for q in questions:
        opts = [q.option_a, q.option_b, q.option_c, q.option_d]
        answer_index = {"A": 1, "B": 2, "C": 3, "D": 4}.get(q.correct_answer or "A", 1)
        writer.writerow([q.question_text] + opts + [answer_index, 30, ""])
    return buf.getvalue().encode("utf-8"), f"{name}_quizizz.csv", "text/csv"


def _export_kahoot(questions, name) -> tuple[bytes, str, str]:
    """Kahoot CSV import format."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Question", "Answer 1", "Answer 2", "Answer 3", "Answer 4",
                     "Time Limit", "Correct Answer"])
    for q in questions:
        correct_num = {"A": 1, "B": 2, "C": 3, "D": 4}.get(q.correct_answer or "A", 1)
        writer.writerow([
            q.question_text,
            q.option_a, q.option_b, q.option_c, q.option_d,
            20, correct_num,
        ])
    return buf.getvalue().encode("utf-8"), f"{name}_kahoot.csv", "text/csv"


def _export_excel(questions, name) -> tuple[bytes, str, str]:
    """Export to .xlsx using openpyxl."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return _export_csv(questions, name)[0], f"{name}.csv", "text/csv"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Questions"

    headers = ["#", "Question", "A", "B", "C", "D", "Answer", "Marks", "Topic", "Difficulty", "Explanation"]
    ws.append(headers)

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="6D28D9")
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for i, q in enumerate(questions, 1):
        ws.append([
            i, q.question_text,
            q.option_a, q.option_b, q.option_c, q.option_d,
            q.correct_answer or "",
            float(q.marks),
            q.topic or "",
            q.difficulty or "",
            q.explanation or "",
        ])

    ws.column_dimensions["B"].width = 60
    for col in ["C", "D", "E", "F"]:
        ws.column_dimensions[col].width = 30

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), f"{name}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _export_pdf(questions, name) -> tuple[bytes, str, str]:
    """Export to PDF using weasyprint (HTML → PDF)."""
    try:
        from weasyprint import HTML as WP_HTML
    except ImportError:
        return _export_txt(questions, name)[0], f"{name}.txt", "text/plain"

    html_parts = [
        "<html><head><meta charset='utf-8'>",
        "<style>",
        "body{font-family:Arial,sans-serif;font-size:12pt;padding:20px}",
        "h1{color:#6D28D9;margin-bottom:24px}",
        ".q{margin-bottom:20px;page-break-inside:avoid}",
        ".q-num{font-weight:bold;color:#6D28D9}",
        ".q-text{margin:4px 0 8px;font-size:13pt}",
        ".opts{display:grid;grid-template-columns:1fr 1fr;gap:4px;margin-bottom:6px}",
        ".opt{padding:4px 8px;border:1px solid #ddd;border-radius:4px}",
        ".opt.correct{background:#d1fae5;border-color:#10b981;font-weight:bold}",
        ".ans{font-size:10pt;color:#6D28D9;margin-top:4px}",
        ".exp{font-size:10pt;color:#666;margin-top:4px;font-style:italic}",
        "</style></head><body>",
        f"<h1>{name}</h1>",
    ]

    for i, q in enumerate(questions, 1):
        html_parts.append(f"<div class='q'>")
        html_parts.append(f"<span class='q-num'>Q{i}.</span>")
        html_parts.append(f"<div class='q-text'>{q.question_text}</div>")
        html_parts.append("<div class='opts'>")
        for letter, text in [("A", q.option_a), ("B", q.option_b), ("C", q.option_c), ("D", q.option_d)]:
            if text:
                css = "opt correct" if q.correct_answer == letter else "opt"
                html_parts.append(f"<div class='{css}'><b>{letter}.</b> {text}</div>")
        html_parts.append("</div>")
        if q.correct_answer:
            html_parts.append(f"<div class='ans'>Answer: {q.correct_answer}</div>")
        if q.explanation:
            html_parts.append(f"<div class='exp'>Explanation: {q.explanation}</div>")
        if q.topic or q.difficulty:
            meta = " | ".join(x for x in [q.topic, q.difficulty] if x)
            html_parts.append(f"<div class='ans' style='color:#888'>{meta}</div>")
        html_parts.append("</div>")

    html_parts.append("</body></html>")
    html_str = "".join(html_parts)

    try:
        buf = io.BytesIO()
        WP_HTML(string=html_str).write_pdf(buf)
        return buf.getvalue(), f"{name}.pdf", "application/pdf"
    except Exception as e:
        # WeasyPrint can fail at render time (not just import) when its native
        # deps — Pango/Cairo — or an incompatible pydyf are missing on the host.
        # Degrade to a plain-text export rather than 500-ing the whole request.
        logger.warning("PDF render failed, falling back to TXT: %s", e)
        return _export_txt(questions, name)[0], f"{name}.txt", "text/plain"


def _export_docx(questions, name) -> tuple[bytes, str, str]:
    """Export to .docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        return _export_txt(questions, name)[0], f"{name}.txt", "text/plain"

    doc = Document()
    doc.add_heading(name, level=1)

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{i}. ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
        p.add_run(q.question_text)

        for letter, text in [("A", q.option_a), ("B", q.option_b), ("C", q.option_c), ("D", q.option_d)]:
            if text:
                op = doc.add_paragraph(f"    {letter}. {text}", style="List Bullet")
                if q.correct_answer == letter:
                    op.runs[0].bold = True

        if q.correct_answer:
            ans_p = doc.add_paragraph()
            run = ans_p.add_run(f"Answer: {q.correct_answer}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

        if q.explanation:
            exp_p = doc.add_paragraph(f"Explanation: {q.explanation}")
            exp_p.runs[0].italic = True

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), f"{name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
