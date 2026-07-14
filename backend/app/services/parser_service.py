"""
parser_service.py — Top-level parsing orchestrator.

Handles:
  - Text parsing (synchronous, fast)
  - File parsing (PDF, DOCX, TXT, CSV)
  - OCR routing (Tesseract → Google Vision fallback)
  - AI cleanup for broken questions
  - Job status updates
"""

import io
import logging
import os
import tempfile
import time
import uuid
from pathlib import Path
from typing import Optional

from app.config import settings
from app.schemas.parser import ParsedQuestion, ParseOptions, ParseResult
from app.parser.text_parser import text_parser
from app.services.ai_service import ai_service

logger = logging.getLogger(__name__)


# ─── File text extractors ─────────────────────────────────────────────────────

def _extract_text_from_pdf(file_bytes: bytes) -> tuple[str, bool]:
    """
    Extract text from PDF. Returns (text, is_ocr_needed).
    Tries pdfplumber first (text layer), falls back to fitz.
    """
    import pdfplumber

    text_parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        logger.warning("pdfplumber failed: %s — trying pymupdf", e)
        try:
            import fitz  # pymupdf
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text_parts.append(page.get_text())
        except Exception as fe:
            logger.error("pymupdf also failed: %s", fe)

    full_text = "\n\n".join(t for t in text_parts if t.strip())
    is_ocr_needed = len(full_text.strip()) < 50  # less than 50 chars → likely scanned

    return full_text, is_ocr_needed


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file with encoding detection."""
    import chardet

    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return file_bytes.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        return file_bytes.decode("utf-8", errors="replace")


def _extract_text_from_csv(file_bytes: bytes) -> str:
    """Extract MCQ data from CSV (question, optA, optB, optC, optD, answer)."""
    import csv

    text = _extract_text_from_txt(file_bytes)
    lines: list[str] = []
    reader = csv.reader(io.StringIO(text))
    headers = None

    for i, row in enumerate(reader):
        if i == 0:
            headers = [h.lower().strip() for h in row]
            continue
        if len(row) >= 5:
            lines.append(f"{i}. {row[0]}")
            lines.append(f"A. {row[1]}")
            lines.append(f"B. {row[2]}")
            lines.append(f"C. {row[3]}")
            lines.append(f"D. {row[4]}")
            if len(row) >= 6 and row[5].strip():
                lines.append(f"Answer: {row[5].strip()}")
            lines.append("")

    return "\n".join(lines) if lines else text


def _ocr_image_bytes(image_bytes: bytes, mime_type: str = "image/jpeg") -> str:
    """Run Tesseract OCR on image bytes."""
    import pytesseract
    from PIL import Image

    if settings.TESSERACT_CMD != "tesseract":
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    image = Image.open(io.BytesIO(image_bytes))

    # Preprocessing for better OCR accuracy
    image = _preprocess_image(image)

    text = pytesseract.image_to_string(
        image,
        lang=settings.OCR_LANGUAGES,
        config="--psm 6 --oem 1",  # PSM 6: assume uniform block, OEM 1: LSTM only
    )
    return text


def _preprocess_image(image):
    """Convert to grayscale and increase contrast for better OCR."""
    import cv2
    import numpy as np
    from PIL import Image

    # Convert PIL to OpenCV
    cv_img = cv2.cvtColor(np.array(image.convert("RGB")), cv2.COLOR_RGB2BGR)

    # Grayscale
    gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)

    # Denoise
    denoised = cv2.fastNlMeansDenoising(gray, h=10)

    # Adaptive threshold (handles uneven lighting). Block size must scale with
    # image resolution — a fixed small window (e.g. 11px) is narrower than a
    # single character stroke on higher-DPI renders (300 DPI scans, PDF pages
    # rasterized for OCR), which erodes strokes instead of correcting lighting
    # and can wipe out entire lines of text.
    block_size = max(11, (min(denoised.shape[:2]) // 30) | 1)
    thresh = cv2.adaptiveThreshold(
        denoised, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, block_size, 2
    )

    return Image.fromarray(thresh)


def _ocr_pdf_as_images(file_bytes: bytes) -> str:
    """Convert PDF pages to images and OCR each one."""
    try:
        import fitz  # pymupdf
        from PIL import Image

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        all_text: list[str] = []

        for page in doc:
            # Render at 300 DPI
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("jpeg")
            page_text = _ocr_image_bytes(img_bytes, "image/jpeg")
            if page_text.strip():
                all_text.append(page_text)

        return "\n\n".join(all_text)
    except Exception as e:
        logger.error("PDF OCR failed: %s", e)
        return ""


# ─── Main Parser Service ──────────────────────────────────────────────────────

class ParserService:

    def parse_text(
        self,
        raw_text: str,
        options: ParseOptions,
        job_id: Optional[str] = None,
    ) -> ParseResult:
        """Parse pasted text directly."""
        job_id = job_id or str(uuid.uuid4())
        result = text_parser.parse(raw_text, options, job_id)

        # AI cleanup for invalid questions
        if settings.AI_CLEANUP_ENABLED and options.ai_cleanup:
            result = self._apply_ai_cleanup(raw_text, result, options)

        return result

    def parse_file(
        self,
        file_bytes: bytes,
        filename: str,
        options: ParseOptions,
        job_id: Optional[str] = None,
    ) -> ParseResult:
        """Parse an uploaded file (PDF, DOCX, TXT, image)."""
        job_id = job_id or str(uuid.uuid4())
        ext = Path(filename).suffix.lower()

        raw_text = ""
        ocr_used = False

        try:
            if ext == ".pdf":
                raw_text, needs_ocr = _extract_text_from_pdf(file_bytes)
                if needs_ocr:
                    logger.info("PDF appears scanned — running OCR")
                    raw_text = _ocr_pdf_as_images(file_bytes)
                    ocr_used = True

            elif ext == ".docx":
                raw_text = _extract_text_from_docx(file_bytes)

            elif ext == ".txt":
                raw_text = _extract_text_from_txt(file_bytes)

            elif ext == ".csv":
                raw_text = _extract_text_from_csv(file_bytes)

            elif ext in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"):
                raw_text = _ocr_image_bytes(file_bytes)
                ocr_used = True

            else:
                # Attempt plain text
                raw_text = _extract_text_from_txt(file_bytes)

        except Exception as e:
            logger.error("File extraction failed for %s: %s", filename, e)
            return ParseResult(
                job_id=job_id,
                questions=[],
                warnings=[f"File extraction failed: {str(e)}"],
                duplicates_removed=0,
                parse_duration_ms=0,
                format_detected=None,
                total_found=0,
            )

        if not raw_text.strip():
            return ParseResult(
                job_id=job_id,
                questions=[],
                warnings=["No text could be extracted from the file"],
                duplicates_removed=0,
                parse_duration_ms=0,
                format_detected=None,
                total_found=0,
            )

        result = text_parser.parse(raw_text, options, job_id)

        if ocr_used:
            result.warnings.insert(0, f"OCR was used to extract text from {ext.upper()}")

        if settings.AI_CLEANUP_ENABLED and options.ai_cleanup:
            result = self._apply_ai_cleanup(raw_text, result, options)

        return result

    # ─── AI cleanup ──────────────────────────────────────────────────────────

    def _apply_ai_cleanup(
        self,
        raw_text: str,
        result: ParseResult,
        options: ParseOptions,
    ) -> ParseResult:
        """
        For questions with validation errors, attempt AI repair.
        Limits to first 10 invalid questions to control API cost.
        """
        invalid = [(q, i) for i, q in enumerate(result.questions) if not q.is_valid]
        if not invalid:
            return result

        # Cap to avoid excessive API cost
        to_repair = invalid[:10]
        repaired_count = 0

        for q, idx in to_repair:
            # Extract the original raw block for this question
            raw_block = self._find_raw_block(raw_text, q)
            if raw_block:
                repaired = ai_service.cleanup_question(raw_block, q)
                if repaired:
                    result.questions[idx] = repaired
                    repaired_count += 1

        if repaired_count:
            result.warnings.append(f"AI cleanup repaired {repaired_count} question(s)")

        # Re-validate after AI cleanup
        from app.parser.validator import validator
        result.questions, new_warnings, dups = validator.validate_all(
            result.questions, remove_duplicates=False
        )

        return result

    def _find_raw_block(self, full_text: str, question: ParsedQuestion) -> Optional[str]:
        """Attempt to find the raw text block for a question by its text."""
        if not question.question_text:
            return None

        # Look for question text in the raw source
        q_snippet = question.question_text[:40].strip()
        idx = full_text.find(q_snippet)
        if idx == -1:
            return None

        # Extract ~500 chars around it
        start = max(0, idx - 20)
        end = min(len(full_text), idx + 500)
        return full_text[start:end]


# ─── Singleton ────────────────────────────────────────────────────────────────
parser_service = ParserService()
